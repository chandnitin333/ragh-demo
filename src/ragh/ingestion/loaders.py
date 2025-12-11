from pathlib import Path
import tempfile
import io
from typing import Optional
import pdfplumber
import docx
from PIL import Image
import pytesseract
import speech_recognition as sr
import moviepy.editor as mpy
from loguru import logger

def load_pdf(path: Path) -> str:
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text.append(page_text)
    logger.debug(f"Loaded PDF {path} length={sum(len(p) for p in text)}")
    return "\n".join(text)

def load_docx(path: Path) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def load_generic(path: Path) -> str:
    parsed = parser.from_file(str(path))
    return parsed.get("content", "") or ""



def extract_text_from_bytes(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_bytes(data)
    if ext in (".docx", ".doc"):
        return extract_docx_bytes(data)
    if ext in (".txt",):
        return data.decode("utf-8", errors="ignore")
    if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
        return extract_image_bytes(data)
    if ext in (".mp3", ".wav", ".m4a"):
        return extract_audio_bytes(data)
    if ext in (".mp4", ".mov", ".mkv", ".avi"):
        return extract_video_bytes(data)
    # fallback: try decode
    return data.decode("utf-8", errors="ignore")

def extract_pdf_bytes(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".pdf") as tmp:
        tmp.write(data); tmp.flush()
        text_pages = []
        with pdfplumber.open(tmp.name) as pdf:
            for p in pdf.pages:
                page_text = p.extract_text() or ""
                text_pages.append(page_text)
        return "\n\n".join(text_pages)

def extract_docx_bytes(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
        tmp.write(data); tmp.flush()
        doc = docx.Document(tmp.name)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paras)

def extract_image_bytes(data: bytes) -> str:
    img = Image.open(io.BytesIO(data))
    text = pytesseract.image_to_string(img)
    return text

def extract_audio_bytes(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".wav") as tmp:
        tmp.write(data); tmp.flush()
        recognizer = sr.Recognizer()
        with sr.AudioFile(tmp.name) as source:
            audio = recognizer.record(source)
            # uses Google STT (requires internet) — replace with whisper call if preferred
            text = recognizer.recognize_google(audio)
            return text

def extract_video_bytes(data: bytes) -> str:
    # Extract audio track and transcribe
    with tempfile.NamedTemporaryFile(suffix=".mp4") as tmp:
        tmp.write(data); tmp.flush()
        clip = mpy.VideoFileClip(tmp.name)
        audio_path = tmp.name + ".wav"
        clip.audio.write_audiofile(audio_path, logger=None)
        # use speech_recognition to read wav
        recognizer = sr.Recognizer()
        with sr.AudioFile(audio_path) as source:
            audio = recognizer.record(source)
            text = recognizer.recognize_google(audio)
            return text